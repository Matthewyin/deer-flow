"""knowledge_upload tool: ingest documents into the ops-knowledge RAG."""

import hashlib
import logging
import uuid
from pathlib import Path
from typing import Optional

from config import get_config
from db.metadata_client import MetadataClient
from rag.ops_knowledge_rag import OpsKnowledgeRAG, chunk_by_doc_type

logger = logging.getLogger(__name__)

_rag: Optional[OpsKnowledgeRAG] = None
_db: Optional[MetadataClient] = None


def _get_rag() -> OpsKnowledgeRAG:
    global _rag
    if _rag is None:
        cfg = get_config()
        _rag = OpsKnowledgeRAG(
            persist_dir=cfg.chroma.persist_dir,
            ollama_base_url=cfg.chroma.ollama_base_url,
            ollama_model=cfg.chroma.ollama_model,
            collection_name=cfg.chroma.collection_name,
        )
        _rag.initialize()
    return _rag


def _get_db() -> MetadataClient:
    global _db
    if _db is None:
        cfg = get_config()
        _db = MetadataClient(db_path=cfg.sqlite.db_path)
    return _db


def _file_hash(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()


def _read_file(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext in (".md", ".txt", ".csv"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == ".pdf":
        return _read_pdf(file_path)
    elif ext == ".docx":
        return _read_docx(file_path)
    elif ext == ".xlsx":
        return _read_xlsx(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def _read_pdf(file_path: str) -> str:
    try:
        import fitz

        doc = fitz.open(file_path)
        return "\n".join(page.get_text() for page in doc)
    except ImportError:
        logger.warning("PyMuPDF not installed, reading PDF as plain text")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.warning(f"Failed to read PDF {file_path}: {e}")
        return ""


def _read_docx(file_path: str) -> str:
    try:
        from docx import Document

        doc = Document(file_path)
        parts: list[str] = []

        # Interleave paragraphs and tables in document order
        # python-docx tracks both via doc.element.body which contains
        # <w:p> (paragraph) and <w:tbl> (table) elements in order.
        body = doc.element.body
        para_idx = 0
        table_idx = 0

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                if para_idx < len(doc.paragraphs):
                    text = doc.paragraphs[para_idx].text.strip()
                    if text:
                        parts.append(text)
                para_idx += 1
            elif tag == "tbl":
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    rows_text: list[str] = []
                    for row in table.rows:
                        cells = [
                            cell.text.strip().replace("\n", " ") for cell in row.cells
                        ]
                        if any(cells):
                            rows_text.append(" | ".join(cells))
                    if rows_text:
                        parts.append("\n".join(rows_text))
                table_idx += 1

        return "\n\n".join(parts)
    except ImportError:
        logger.warning("python-docx not installed, reading DOCX as plain text")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.warning(f"Failed to read DOCX {file_path}: {e}")
        return ""


def _read_xlsx(file_path: str) -> str:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append("\t".join(cells))
        return "\n".join(rows)
    except ImportError:
        logger.warning("openpyxl not installed, reading XLSX as plain text")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.warning(f"Failed to read XLSX {file_path}: {e}")
        return ""


def register(mcp):
    @mcp.tool()
    def knowledge_upload(
        file_path: str,
        doc_type: str,
        title: str,
        device_vendor: Optional[str] = None,
        device_type: Optional[str] = None,
    ) -> dict:
        """
        将运维文档入库到知识库。支持 pdf/docx/xlsx/txt/md/csv 格式。
        自动去重（同文件hash跳过，同标题覆盖更新）。

        Args:
            file_path: 文件路径（相对于raw_dir或绝对路径）
            doc_type: 文档类型，可选值: fault, sop, emergency, solution, event
            title: 文档标题
            device_vendor: 设备厂商（可选），如 huawei, h3c, hillstone, f5
            device_type: 设备类型（可选），如 router, switch, firewall, load_balancer

        Returns:
            dict: 入库结果，包含 doc_id, chunk_count, status
        """
        rag = _get_rag()
        db = _get_db()

        cfg = get_config()
        raw_dir = cfg.chroma.raw_dir
        project_root = Path(__file__).resolve().parent.parent.parent.parent

        resolved = Path(file_path)
        if not resolved.is_absolute():
            resolved = project_root / file_path

        if not resolved.exists():
            return {"status": "error", "message": f"文件不存在: {file_path}"}

        file_hash = _file_hash(str(resolved))
        existing = db.get_document_by_hash(file_hash)
        if existing:
            db.log_upload(
                existing["doc_id"], "upload", "skipped", "文件hash未变更，跳过入库"
            )
            return {
                "status": "skipped",
                "doc_id": existing["doc_id"],
                "message": "文件内容未变更，跳过入库",
            }

        existing_by_title = db.get_document_by_title(title)
        if existing_by_title:
            old_doc_id = existing_by_title["doc_id"]
            deleted = rag.delete_chunks_by_doc_id(old_doc_id)
            db.delete_document(old_doc_id)
            db.log_upload(
                old_doc_id,
                "overwrite",
                "success",
                f"标题重复，删除旧版本({deleted}个chunks)",
            )
            logger.info(
                f"Overwriting doc '{title}' (old_id={old_doc_id}, deleted={deleted} chunks)"
            )

        content = _read_file(str(resolved))
        if not content.strip():
            return {"status": "error", "message": "文件内容为空"}

        raw_chunks = chunk_by_doc_type(content, doc_type)
        if not raw_chunks:
            return {"status": "error", "message": "分块后无有效内容"}

        doc_id = str(uuid.uuid4())
        import datetime

        upload_date = datetime.datetime.now().isoformat()

        source_rel = (
            str(resolved.relative_to(project_root))
            if str(resolved).startswith(str(project_root))
            else str(resolved)
        )

        chunks = []
        for idx, chunk_content in enumerate(raw_chunks):
            chunks.append(
                {
                    "content": chunk_content,
                    "metadata": {
                        "doc_id": doc_id,
                        "doc_type": doc_type,
                        "device_vendor": device_vendor or "",
                        "device_type": device_type or "",
                        "source_file": source_rel,
                        "chunk_index": idx,
                        "total_chunks": len(raw_chunks),
                        "upload_date": upload_date,
                    },
                }
            )

        added = rag.add_chunks(chunks)

        doc_record = {
            "doc_id": doc_id,
            "doc_type": doc_type,
            "title": title,
            "source_file": source_rel,
            "device_vendor": device_vendor,
            "device_type": device_type,
            "chunk_count": added,
            "upload_date": upload_date,
            "file_hash": file_hash,
        }
        db.upsert_document(doc_record)
        db.log_upload(doc_id, "upload", "success", f"入库{added}个chunks")

        return {
            "status": "success",
            "doc_id": doc_id,
            "title": title,
            "doc_type": doc_type,
            "chunk_count": added,
            "file_hash": file_hash[:16],
        }
