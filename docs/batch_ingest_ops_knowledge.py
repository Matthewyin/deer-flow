#!/usr/bin/env python3
"""Batch ingest all documents from docs/ops-knowledge/raw/ into ops_knowledge RAG.

Directory mapping:
  raw/01 系统应急预案/* → doc_type=emergency_system
  raw/02 网络应急预案/* → doc_type=emergency_network
  raw/03 安全应急预案/* → doc_type=emergency_security
  raw/SOP/*            → doc_type=sop

Usage:
  python docs/batch_ingest_ops_knowledge.py [--dry-run] [--limit N] [--retry RETRIES]
"""

import hashlib
import sys
import uuid
import time
import signal
import datetime
import argparse
from pathlib import Path

# Add mcp-servers to path
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
MCP_SERVERS = PROJECT_ROOT / "mcp-servers" / "ops-knowledge"
sys.path.insert(0, str(MCP_SERVERS))

from config import get_config
from db.metadata_client import MetadataClient
from rag.ops_knowledge_rag import OpsKnowledgeRAG, chunk_by_doc_type

# --- File readers (same logic as knowledge_upload tool) ---


def _file_hash(file_path: str) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()


def _read_file(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext in (".md", ".txt", ".csv"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == ".pdf":
        return _read_pdf(file_path)
    elif ext == ".docx":
        return _read_docx(file_path)
    elif ext == ".xlsx":
        return _read_xlsx(file_path)
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def _read_pdf(file_path: str) -> str:
    try:
        import fitz

        doc = fitz.open(file_path)
        return "\n".join(page.get_text() for page in doc)
    except ImportError:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  [WARN] Failed to read PDF {file_path}: {e}")
        return ""


def _read_docx(file_path: str) -> str:
    try:
        from docx import Document

        doc = Document(file_path)
        parts: list[str] = []

        # Interleave paragraphs and tables in document order
        # python-docx tracks both via doc.element.body which contains
        # <w:p> (paragraph) and <w:tbl> (table) elements in order.
        body = doc.element.body
        para_idx = 0
        table_idx = 0

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                if para_idx < len(doc.paragraphs):
                    text = doc.paragraphs[para_idx].text.strip()
                    if text:
                        parts.append(text)
                para_idx += 1
            elif tag == "tbl":
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    rows_text: list[str] = []
                    for row in table.rows:
                        cells = [
                            cell.text.strip().replace("\n", " ") for cell in row.cells
                        ]
                        if any(cells):
                            rows_text.append(" | ".join(cells))
                    if rows_text:
                        parts.append("\n".join(rows_text))
                table_idx += 1

        return "\n\n".join(parts)
    except ImportError:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  [WARN] Failed to read DOCX {file_path}: {e}")
        return ""


def _read_xlsx(file_path: str) -> str:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append("\t".join(cells))
        return "\n".join(rows)
    except ImportError:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"  [WARN] Failed to read XLSX {file_path}: {e}")
        return ""


_DIR_MAP = {
    "01 系统应急预案": "emergency_system",
    "02 网络应急预案": "emergency_network",
    "03 安全应急预案": "emergency_security",
    "SOP": "sop",
}


def classify_doc_type(rel_path: str) -> str:
    parts = Path(rel_path).parts
    # parts e.g. ("docs", "ops-knowledge", "raw", "01 系统应急预案", "file.docx")
    for part in parts:
        if part in _DIR_MAP:
            return _DIR_MAP[part]
    return "sop"


def extract_title(filename: str) -> str:
    """Extract a clean title from filename."""
    stem = Path(filename).stem
    # Remove common prefixes like IT0001- or AQ001-
    # Just use the full stem as title, it's descriptive enough
    return stem.strip()


class EmbeddingTimeout(Exception):
    pass


def _timeout_handler(signum, frame):
    raise EmbeddingTimeout("Embedding operation timed out")


def ingest_one_file(
    rag, db, f, raw_dir, project_root, max_retries=3, timeout_per_file=300
):
    """Ingest a single file with timeout and retry logic.

    Returns: (status, chunks_added)
        status: 'ok', 'skip_hash', 'skip_empty', 'error', 'timeout'
    """
    rel = f.relative_to(raw_dir)
    doc_type = classify_doc_type(str(f.relative_to(project_root)))
    title = extract_title(f.name)
    n_chunks = 0

    # Check hash dedup
    try:
        file_hash = _file_hash(str(f))
    except Exception as e:
        print(f"  [SKIP] Cannot read file hash: {e}")
        return "error", 0

    existing = db.get_document_by_hash(file_hash)
    if existing:
        print(f"  [SKIP] Hash unchanged (doc_id={existing['doc_id'][:8]}...)")
        return "skip_hash", 0

    # Check title dedup
    existing_by_title = db.get_document_by_title(title)
    if existing_by_title:
        old_doc_id = existing_by_title["doc_id"]
        deleted = rag.delete_chunks_by_doc_id(old_doc_id)
        db.delete_document(old_doc_id)
        print(f"  [OVERWRITE] Title exists, deleted {deleted} old chunks")

    # Read file
    content = _read_file(str(f))
    if not content.strip():
        print(f"  [SKIP] Empty content")
        return "skip_empty", 0

    # Chunk
    raw_chunks = chunk_by_doc_type(content, doc_type)
    if not raw_chunks:
        print(f"  [SKIP] No valid chunks after splitting")
        return "skip_empty", 0

    # Build chunks with metadata
    doc_id = str(uuid.uuid4())
    upload_date = datetime.datetime.now().isoformat()
    source_rel = str(f.relative_to(project_root))

    chunks = []
    for idx, chunk_content in enumerate(raw_chunks):
        chunks.append(
            {
                "content": chunk_content,
                "metadata": {
                    "doc_id": doc_id,
                    "doc_type": doc_type,
                    "device_vendor": "",
                    "device_type": "",
                    "source_file": source_rel,
                    "chunk_index": idx,
                    "total_chunks": len(raw_chunks),
                    "upload_date": upload_date,
                },
            }
        )

    # Add to vectorstore in mini-batches to avoid Ollama overload
    MINI_BATCH_SIZE = 20
    total_added = 0
    failed = False

    for batch_start in range(0, len(chunks), MINI_BATCH_SIZE):
        batch = chunks[batch_start : batch_start + MINI_BATCH_SIZE]
        batch_num = batch_start // MINI_BATCH_SIZE + 1
        total_batches = (len(chunks) + MINI_BATCH_SIZE - 1) // MINI_BATCH_SIZE

        last_error = None
        success = False
        for attempt in range(1, max_retries + 1):
            try:
                old_handler = signal.signal(signal.SIGALRM, _timeout_handler)
                signal.alarm(120)  # 120s per mini-batch
                try:
                    n = rag.add_chunks(batch)
                finally:
                    signal.alarm(0)
                    signal.signal(signal.SIGALRM, old_handler)
                total_added += n
                success = True
                break
            except EmbeddingTimeout:
                signal.alarm(0)
                print(
                    f"  [TIMEOUT] Batch {batch_num}/{total_batches} attempt {attempt}/{max_retries}"
                )
                last_error = "timeout"
                if attempt < max_retries:
                    time.sleep(5)
            except Exception as e:
                signal.alarm(0)
                print(
                    f"  [ERROR] Batch {batch_num}/{total_batches} attempt {attempt}: {e}"
                )
                last_error = str(e)
                if attempt < max_retries:
                    time.sleep(5)

        if not success:
            print(f"  [FAILED] Batch {batch_num}/{total_batches} failed: {last_error}")
            failed = True
            break

        if batch_num % 5 == 0 or batch_num == total_batches:
            print(f"  [PROGRESS] {total_added}/{len(chunks)} chunks embedded")

        # Small delay between mini-batches
        time.sleep(0.5)

    if failed:
        # Clean up partial data
        try:
            rag.delete_chunks_by_doc_id(doc_id)
        except Exception:
            pass
        return "error", 0

    # Record in DB
    doc_record = {
        "doc_id": doc_id,
        "doc_type": doc_type,
        "title": title,
        "source_file": source_rel,
        "device_vendor": "",
        "device_type": "",
        "chunk_count": total_added,
        "upload_date": upload_date,
        "file_hash": file_hash,
    }
    db.upsert_document(doc_record)
    db.log_upload(doc_id, "upload", "success", f"Batch ingest: {total_added} chunks")
    print(f"  [OK] {total_added} chunks ingested (doc_id={doc_id[:8]}...)")
    return "ok", total_added


def main():
    parser = argparse.ArgumentParser(description="Batch ingest ops-knowledge docs")
    parser.add_argument(
        "--dry-run", action="store_true", help="Only list files, don't ingest"
    )
    parser.add_argument(
        "--limit", type=int, default=0, help="Max files to ingest (0=all)"
    )
    parser.add_argument(
        "--retry", type=int, default=3, help="Max retries per file (default=3)"
    )
    parser.add_argument(
        "--timeout",
        type=int,
        default=300,
        help="Timeout per file in seconds (default=300)",
    )
    parser.add_argument(
        "--delay",
        type=float,
        default=2.0,
        help="Delay between files in seconds (default=2.0)",
    )
    args = parser.parse_args()

    cfg = get_config()
    raw_dir = (PROJECT_ROOT / cfg.chroma.raw_dir).resolve()

    if not raw_dir.exists():
        print(f"ERROR: raw_dir not found: {raw_dir}")
        sys.exit(1)

    # Collect all supported files
    supported_ext = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".csv"}
    files = []
    for f in sorted(raw_dir.rglob("*")):
        if f.is_file() and f.suffix.lower() in supported_ext:
            files.append(f)

    print(f"Found {len(files)} files in {raw_dir}")

    if args.dry_run:
        for f in files:
            rel = f.relative_to(raw_dir)
            doc_type = classify_doc_type(str(f.relative_to(PROJECT_ROOT)))
            title = extract_title(f.name)
            print(f"  [{doc_type:9s}] {title}")
        print(f"\nTotal: {len(files)} files (dry-run, no ingestion)")
        return

    # Initialize RAG and DB
    print("\nInitializing RAG...")
    rag = OpsKnowledgeRAG(
        persist_dir=cfg.chroma.persist_dir,
        ollama_base_url=cfg.chroma.ollama_base_url,
        ollama_model=cfg.chroma.ollama_model,
        collection_name=cfg.chroma.collection_name,
    )
    rag.initialize()
    db = MetadataClient(db_path=cfg.sqlite.db_path)

    current_count = rag._vectorstore._collection.count()
    print(f"Current vectorstore chunks: {current_count}")

    # Stats
    stats = {"ok": 0, "skip_hash": 0, "skip_empty": 0, "error": 0, "total_chunks": 0}
    start_time = time.time()

    for i, f in enumerate(files, 1):
        if args.limit > 0 and stats["ok"] >= args.limit:
            print(f"\n[LIMIT] Reached --limit={args.limit}, stopping.")
            break

        rel = f.relative_to(raw_dir)
        doc_type = classify_doc_type(str(f.relative_to(PROJECT_ROOT)))
        title = extract_title(f.name)
        elapsed = int(time.time() - start_time)

        print(f"\n[{i}/{len(files)}] {rel} ({doc_type}) [elapsed: {elapsed}s]")

        status, n = ingest_one_file(
            rag,
            db,
            f,
            raw_dir,
            PROJECT_ROOT,
            max_retries=args.retry,
            timeout_per_file=args.timeout,
        )
        stats[status] += 1
        stats["total_chunks"] += n

        # Delay between files to avoid overwhelming Ollama
        if i < len(files) and args.delay > 0:
            time.sleep(args.delay)

    # Summary
    total_time = int(time.time() - start_time)
    final_count = rag._vectorstore._collection.count()
    print(f"\n{'=' * 60}")
    print(f"Batch ingest complete (took {total_time}s):")
    print(f"  Success:     {stats['ok']}")
    print(f"  Skipped:     {stats['skip_hash']} (hash unchanged)")
    print(f"  Empty:       {stats['skip_empty']} (no content)")
    print(f"  Errors:      {stats['error']}")
    print(f"  New chunks:  {stats['total_chunks']}")
    print(f"  Final count: {final_count} total chunks in vectorstore")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
